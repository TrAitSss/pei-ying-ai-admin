import os
import json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional
from jinja2 import Template as JinjaTemplate
from docx import Document as DocxDocument
from docx.shared import Pt
import tempfile

from app.database import get_db
from app.models import Template, GeneratedDocument, User
from app.schemas import TemplateCreate, TemplateResponse, GenerateDocumentRequest
from app.routers.auth import get_current_user

router = APIRouter()

@router.post("/", response_model=TemplateResponse)
async def create_template(
    template: TemplateCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    db_template = Template(
        **template.dict(),
        created_by=current_user.id,
    )
    db.add(db_template)
    await db.commit()
    await db.refresh(db_template)
    return db_template

@router.get("/", response_model=List[TemplateResponse])
async def list_templates(
    template_type: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = select(Template).where(Template.is_active == True)
    if template_type:
        query = query.where(Template.template_type == template_type)
    result = await db.execute(query.order_by(Template.created_at.desc()))
    return result.scalars().all()

@router.get("/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    return template

@router.post("/{template_id}/generate")
async def generate_document(
    template_id: int,
    request: GenerateDocumentRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 使用 Jinja2 渲染
    jinja_template = JinjaTemplate(template.content or "")
    rendered_content = jinja_template.render(**request.variables_data)
    
    # 生成 Word 文档
    doc = DocxDocument()
    for line in rendered_content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    output_dir = "./generated"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{request.title}.docx")
    doc.save(output_path)
    
    gen_doc = GeneratedDocument(
        template_id=template_id,
        title=request.title,
        variables_data=request.variables_data,
        file_path=output_path,
        status="draft",
        created_by=current_user.id,
    )
    db.add(gen_doc)
    await db.commit()
    await db.refresh(gen_doc)
    
    return {
        "id": gen_doc.id,
        "title": gen_doc.title,
        "status": gen_doc.status,
        "file_path": output_path,
    }

@router.get("/{template_id}/preview")
async def preview_template(
    template_id: int,
    variables: str = "{}",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    vars_dict = json.loads(variables)
    jinja_template = JinjaTemplate(template.content or "")
    rendered = jinja_template.render(**vars_dict)
    
    return {"preview": rendered, "variables_used": list(vars_dict.keys())}
