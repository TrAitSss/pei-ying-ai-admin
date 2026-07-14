import os
import re
import json
import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional
from jinja2 import Template as JinjaTemplate
from docx import Document as DocxDocument
from docx.shared import Pt, Inches

from app.database import get_db
from app.models import Template, GeneratedDocument, User
from app.schemas import TemplateCreate, TemplateResponse, GenerateDocumentRequest
from app.routers.auth import get_current_user

router = APIRouter()


# ========== 模板 CRUD ==========

@router.post("/", response_model=TemplateResponse)
async def create_template(
    template: TemplateCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    db_template = Template(
        **template.dict(),
        created_by=current_user.id,
    )
    db.add(db_template)
    await db.commit()
    await db.refresh(db_template)
    return db_template


@router.get("/", response_model=List[TemplateResponse])
async def list_templates(
    template_type: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = select(Template).where(Template.is_active == True)
    if template_type:
        query = query.where(Template.template_type == template_type)
    result = await db.execute(query.order_by(Template.created_at.desc()))
    return result.scalars().all()


@router.put("/{template_id}", response_model=TemplateResponse)
async def update_template(
    template_id: int,
    template: TemplateCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    db_template = result.scalar_one_or_none()
    if not db_template:
        raise HTTPException(status_code=404, detail="模板不存在")

    for key, value in template.dict().items():
        setattr(db_template, key, value)

    await db.commit()
    await db.refresh(db_template)
    return db_template


@router.delete("/{template_id}")
async def delete_template(
    template_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    template.is_active = False
    await db.commit()
    return {"message": "模板已删除"}


@router.get("/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    return template


# ========== 上传 .docx 模板 ==========

@router.post("/upload", response_model=TemplateResponse)
async def upload_docx_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    template_type: str = Form("tender"),
    description: str = Form(""),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传 .docx 文件，自动提取 {{ }} 占位符作为变量"""
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="只支持 .docx 格式文件")

    content_bytes = await file.read()
    doc = DocxDocument(io.BytesIO(content_bytes))

    # 提取所有段落文本，保留格式标记
    lines = []
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            lines.append("")
            continue

        # 根据样式添加 markdown 标记
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "list bullet" in style_name:
            lines.append(f"- {text}")
        elif "list number" in style_name:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

    full_content = "\n".join(lines)

    # 正则提取 {{ variable_name }} 占位符
    var_pattern = r'\{\{\s*(\w+)\s*\}\}'
    var_names = list(dict.fromkeys(re.findall(var_pattern, full_content)))

    variables = [{"name": v, "label": v, "default": ""} for v in var_names]

    db_template = Template(
        name=name,
        template_type=template_type,
        description=description or f"从 {file.filename} 导入的模板",
        content=full_content,
        variables=variables,
        is_active=True,
        created_by=current_user.id,
    )
    db.add(db_template)
    await db.commit()
    await db.refresh(db_template)

    return db_template


# ========== 文档生成与预览 ==========

@router.post("/{template_id}/generate")
async def generate_document(
    template_id: int,
    request: GenerateDocumentRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    # Jinja2 渲染
    jinja_template = JinjaTemplate(template.content or "")
    rendered_content = jinja_template.render(**request.variables_data)

    # 生成 Word 文档（支持 markdown 格式）
    doc = _render_to_docx(rendered_content)

    output_dir = "./generated"
    os.makedirs(output_dir, exist_ok=True)
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', request.title)
    output_path = os.path.join(output_dir, f"{safe_title}.docx")
    doc.save(output_path)

    gen_doc = GeneratedDocument(
        template_id=template_id,
        title=request.title,
        variables_data=request.variables_data,
        file_path=output_path,
        status="draft",
        created_by=current_user.id,
    )
    db.add(gen_doc)
    await db.commit()
    await db.refresh(gen_doc)

    return {
        "id": gen_doc.id,
        "title": gen_doc.title,
        "status": gen_doc.status,
        "file_path": output_path,
    }


@router.post("/preview-content")
async def preview_content(
    body: dict,
    current_user: User = Depends(get_current_user),
):
    """实时预览：接收模板内容 + 变量数据，返回渲染结果"""
    content = body.get("content", "")
    variables_data = body.get("variables_data", {})

    try:
        jinja_template = JinjaTemplate(content)
        rendered = jinja_template.render(**variables_data)
    except Exception as e:
        return {"preview": f"[模板语法错误: {str(e)}]", "error": True}

    return {"preview": rendered, "error": False}


@router.get("/{template_id}/preview")
async def preview_template(
    template_id: int,
    variables: str = "{}",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(select(Template).where(Template.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    vars_dict = json.loads(variables)
    jinja_template = JinjaTemplate(template.content or "")
    rendered = jinja_template.render(**vars_dict)

    return {"preview": rendered, "variables_used": list(vars_dict.keys())}


# ========== 已生成文档管理 ==========

@router.get("/documents/list")
async def list_documents(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(GeneratedDocument).order_by(GeneratedDocument.created_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "title": d.title,
            "template_id": d.template_id,
            "status": d.status,
            "created_at": d.created_at.isoformat() if d.created_at else None,
            "download_url": f"/api/templates/documents/{d.id}/download",
        }
        for d in docs
    ]


@router.get("/documents/{doc_id}/download")
async def download_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(GeneratedDocument).where(GeneratedDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    if not doc.file_path or not os.path.isfile(doc.file_path):
        raise HTTPException(status_code=404, detail="文件不存在或已被删除")

    return FileResponse(
        doc.file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{doc.title}.docx",
    )


@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(GeneratedDocument).where(GeneratedDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 删除物理文件
    if doc.file_path and os.path.isfile(doc.file_path):
        os.remove(doc.file_path)

    await db.delete(doc)
    await db.commit()
    return {"message": "文档已删除"}


# ========== 辅助函数：将渲染后的文本转为带格式的 Word 文档 ==========

def _render_to_docx(rendered_content: str) -> DocxDocument:
    """支持 markdown 风格的格式标记：
    # 标题1
    ## 标题2
    ### 标题3
    - 无序列表
    1. 有序列表
    **粗体**  *斜体*
    | 表格 | 行 |
    ---
    分割线
    """
    doc = DocxDocument()
    lines = rendered_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 分割线
        if stripped in ('---', '***', '___'):
            doc.add_paragraph().add_run("_" * 50)
            i += 1
            continue

        # 标题
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_formatted_paragraph(doc, stripped[2:], style='List Bullet')
        # 有序列表
        elif re.match(r'^\d+[\.\)]\s', stripped):
            text = re.sub(r'^\d+[\.\)]\s', '', stripped)
            _add_formatted_paragraph(doc, text, style='List Number')
        # 表格
        elif stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue
        else:
            _add_formatted_paragraph(doc, stripped)

        i += 1

    return doc


def _add_formatted_paragraph(doc, text: str, style: str = None):
    """解析 **粗体** 和 *斜体* 并添加到段落"""
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()

    # 先分割 **粗体**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            _add_italic_runs(para, part[2:-2], bold=True)
        else:
            _add_italic_runs(para, part, bold=False)

    return para


def _add_italic_runs(para, text: str, bold: bool = False):
    """解析 *斜体* 并添加 runs"""
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
            run.bold = bold
        else:
            run = para.add_run(part)
            run.bold = bold


def _add_table(doc, table_lines: list):
    """将 markdown 表格行转为 Word 表格"""
    if not table_lines:
        return

    rows = []
    for line in table_lines:
        # 跳过分割行 |---|---|
        if re.match(r'^\|[\s\-:|]+$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'

    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < len(table.rows[r_idx].cells):
                table.rows[r_idx].cells[c_idx].text = cell_text
